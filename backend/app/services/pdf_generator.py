import io
import re
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.logger import logger

class PDFGenerator:
    def _convert_markdown_to_html(self, text: str) -> str:
        """Converts bold and italic markdown elements to ReportLab HTML tags."""
        # Replace bold markdown (**text** or __text__) with <b>tags
        text = re.sub(r'\*\*(.*?)\*\*|__(.*?)__', r'<b>\1\2</b>', text)
        # Replace italic markdown (*text* or _text_) with <i>tags
        text = re.sub(r'\*(.*?)\*|_(.*?)_', r'<i>\1\2</i>', text)
        return text

    def generate_pdf(self, markdown_text: str, settings: dict = None) -> bytes:
        """
        Render Markdown resume text to a beautiful PDF document using ReportLab.
        Customizable font, accent color, and margins.
        """
        logger.info("Generating PDF resume from optimized markdown...")
        if settings is None:
            settings = {}

        # Parse settings
        font_family = settings.get("font", "Helvetica") # Helvetica, Times-Roman, Courier
        accent_color_hex = settings.get("color", "#2563EB") # Blue accent
        layout_mode = settings.get("layout", "Professional") # Modern, Minimalist, Compact
        
        # Adjust vertical spacing according to mode
        spacer_size = 8
        if layout_mode == "Compact":
            spacer_size = 5
        elif layout_mode == "Minimalist":
            spacer_size = 12

        accent_color = HexColor(accent_color_hex)
        dark_neutral = HexColor("#1E293B") # Slate 800
        light_neutral = HexColor("#475569") # Slate 600

        # Stylesheet setup
        styles = getSampleStyleSheet()
        
        # Modify existing styles to prevent crashes, or add unique ones
        title_style = ParagraphStyle(
            'ResumeTitle',
            parent=styles['Normal'],
            fontName=f"{font_family}-Bold",
            fontSize=22,
            leading=26,
            textColor=accent_color,
            alignment=1, # Center
            spaceAfter=12
        )
        
        h2_style = ParagraphStyle(
            'ResumeH2',
            parent=styles['Normal'],
            fontName=f"{font_family}-Bold",
            fontSize=13,
            leading=16,
            textColor=accent_color,
            spaceBefore=10,
            spaceAfter=4,
            keepWithNext=True
        )

        h3_style = ParagraphStyle(
            'ResumeH3',
            parent=styles['Normal'],
            fontName=f"{font_family}-Bold",
            fontSize=10,
            leading=13,
            textColor=dark_neutral,
            spaceBefore=6,
            spaceAfter=2,
            keepWithNext=True
        )

        body_style = ParagraphStyle(
            'ResumeBody',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=9.5,
            leading=13,
            textColor=dark_neutral,
            spaceAfter=4
        )

        bullet_style = ParagraphStyle(
            'ResumeBullet',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=9.5,
            leading=13,
            textColor=dark_neutral,
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=3
        )

        contact_style = ParagraphStyle(
            'ResumeContact',
            parent=styles['Normal'],
            fontName=font_family,
            fontSize=9,
            leading=12,
            textColor=light_neutral,
            alignment=1, # Center
            spaceAfter=10
        )

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.5*inch,
            leftMargin=0.5*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )

        story = []
        lines = markdown_text.split('\n')
        
        in_list = False
        contact_info_candidate = []
        is_first_h1 = True

        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            # Heading 1 (Candidate Name / Main Title)
            if line_str.startswith('# '):
                in_list = False
                name = line_str[2:].strip()
                story.append(Paragraph(self._convert_markdown_to_html(name), title_style))
                
            # Heading 2 (Main Sections like Experience, Education)
            elif line_str.startswith('## '):
                in_list = False
                title = line_str[3:].strip()
                story.append(Spacer(1, spacer_size))
                story.append(Paragraph(self._convert_markdown_to_html(title), h2_style))
                story.append(HRFlowable(
                    width="100%", 
                    thickness=1, 
                    color=accent_color, 
                    spaceBefore=1, 
                    spaceAfter=spacer_size
                ))

            # Heading 3 (Sub-sections like specific Job Positions or Degrees)
            elif line_str.startswith('### '):
                in_list = False
                subtitle = line_str[4:].strip()
                story.append(Paragraph(self._convert_markdown_to_html(subtitle), h3_style))

            # Bullet List Items
            elif line_str.startswith('- ') or line_str.startswith('* '):
                in_list = True
                bullet_content = line_str[2:].strip()
                bullet_content_html = self._convert_markdown_to_html(bullet_content)
                story.append(Paragraph(f"&bull; {bullet_content_html}", bullet_style))

            # Standard Paragraph
            else:
                # If we are parsing elements right below Name, treat it as Centered Contact Details
                if len(story) == 1 and not line_str.startswith('#') and len(contact_info_candidate) < 3:
                    contact_info_candidate.append(line_str)
                else:
                    in_list = False
                    paragraph_html = self._convert_markdown_to_html(line_str)
                    story.append(Paragraph(paragraph_html, body_style))
                    story.append(Spacer(1, 2))

        # Insert contact details right under Title if found
        if contact_info_candidate:
            contact_str = "  |  ".join(contact_info_candidate)
            # insert contact details at index 1 in the list of flowables (right below title)
            story.insert(1, Paragraph(contact_str, contact_style))
            
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    def generate_docx(self, markdown_text: str, settings: dict = None) -> bytes:
        """
        Render Markdown resume text to a beautiful, clean DOCX document.
        Customizable font, accent color.
        """
        logger.info("Generating DOCX resume from optimized markdown...")
        if settings is None:
            settings = {}

        font_family = settings.get("font", "Arial") # Arial, Calibri, Times New Roman
        accent_color_hex = settings.get("color", "#2563EB") # Blue

        doc = docx.Document()
        
        # Configure standard margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        lines = markdown_text.split('\n')
        is_first_p = True

        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            # Heading 1 (Name)
            if line_str.startswith('# '):
                name = line_str[2:].strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(name)
                run.font.name = font_family
                run.font.size = Pt(22)
                run.bold = True
                # Set color hex
                run.font.color.rgb = docx.shared.RGBColor(
                    int(accent_color_hex[1:3], 16),
                    int(accent_color_hex[3:5], 16),
                    int(accent_color_hex[5:7], 16)
                )
                
            # Heading 2 (Sections)
            elif line_str.startswith('## '):
                title = line_str[3:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(title)
                run.font.name = font_family
                run.font.size = Pt(13)
                run.bold = True
                run.font.color.rgb = docx.shared.RGBColor(
                    int(accent_color_hex[1:3], 16),
                    int(accent_color_hex[3:5], 16),
                    int(accent_color_hex[5:7], 16)
                )
                
            # Heading 3 (Positions / Education details)
            elif line_str.startswith('### '):
                subtitle = line_str[4:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(subtitle)
                run.font.name = font_family
                run.font.size = Pt(10.5)
                run.bold = True
                
            # List items
            elif line_str.startswith('- ') or line_str.startswith('* '):
                bullet_content = line_str[2:].strip()
                # Remove markdown bold tags from text since word uses runs
                cleaned_text = re.sub(r'\*\*(.*?)\*\*|__(.*?)__', r'\1\2', bullet_content)
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(cleaned_text)
                run.font.name = font_family
                run.font.size = Pt(9.5)
                
            # Paragraph
            else:
                # If first block after Heading 1, align center (contact info)
                p = doc.add_paragraph()
                if is_first_p and len(doc.paragraphs) >= 2 and doc.paragraphs[-2].alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(8)
                    is_first_p = False
                else:
                    p.paragraph_format.space_after = Pt(4)
                    
                cleaned_text = re.sub(r'\*\*(.*?)\*\*|__(.*?)__', r'\1\2', line_str)
                run = p.add_run(cleaned_text)
                run.font.name = font_family
                run.font.size = Pt(9.5)

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes

    def generate_pdf_via_latex(self, latex_code: str) -> bytes:
        """
        Compiles LaTeX code to PDF using either a local pdflatex compiler 
        (if available) or the cloud LaTeX.Online compiler API.
        """
        try:
            # First, attempt local compilation if pdflatex is installed
            return self.compile_latex_locally(latex_code)
        except (FileNotFoundError, ValueError) as local_err:
            logger.warning(f"Local LaTeX compilation skipped/failed: {str(local_err)}. Trying cloud compilation...")
            # Fallback to cloud-based compilation
            return self.compile_latex_via_cloud(latex_code)

    def compile_latex_locally(self, latex_code: str) -> bytes:
        """Attempts to compile LaTeX code locally using pdflatex."""
        import subprocess
        import tempfile
        import os
        from shutil import which

        if not which("pdflatex"):
            raise FileNotFoundError("pdflatex executable not found in PATH.")

        logger.info("Compiling LaTeX locally using pdflatex...")
        with tempfile.TemporaryDirectory() as tmpdir:
            tex_path = os.path.join(tmpdir, "resume.tex")
            pdf_path = os.path.join(tmpdir, "resume.pdf")

            with open(tex_path, "w", encoding="utf-8") as f:
                f.write(latex_code)

            # Run pdflatex twice for proper link/page resolution
            for i in range(2):
                logger.info(f"Running pdflatex pass {i+1}...")
                result = subprocess.run(
                    ["pdflatex", "-interaction=nonstopmode", "resume.tex"],
                    cwd=tmpdir,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True
                )
                if result.returncode != 0:
                    logger.error(f"pdflatex failed: {result.stdout[:500]}")
                    raise ValueError(f"pdflatex compilation failed: {result.stdout[:200]}")

            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    return f.read()
            else:
                raise FileNotFoundError("Compiled PDF file not found after running pdflatex.")

    def compile_latex_via_cloud(self, latex_code: str) -> bytes:
        """Compiles LaTeX code using the LaTeX.Online API (via POST file upload)."""
        import requests
        
        logger.info("Compiling LaTeX using cloud LaTeX.Online API (POST file upload)...")
        url = "https://latexonline.cc/compile?target=resume.tex"
        
        try:
            files = {"file": ("resume.tex", latex_code.encode("utf-8"), "text/plain")}
            response = requests.post(url, files=files, timeout=20)
            if response.status_code == 200:
                content = response.content
                # Sanity check: PDF files start with '%PDF'
                if not content.startswith(b'%PDF'):
                    logger.error(f"Cloud LaTeX API returned 200 but content is not a valid PDF. First bytes: {content[:100]}")
                    raise ValueError("Cloud LaTeX API returned a non-PDF response despite HTTP 200.")
                logger.info("LaTeX compiled successfully via cloud API.")
                return content
            else:
                error_log = response.text[:1000]
                logger.error(f"Cloud LaTeX API failed ({response.status_code}): {error_log}")
                raise ValueError(f"Cloud LaTeX compilation failed (HTTP {response.status_code}). Log: {error_log[:300]}")
        except Exception as e:
            logger.error(f"Cloud LaTeX API request failed: {str(e)}")
            raise e

