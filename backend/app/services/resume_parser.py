import io
import pdfplumber
import docx
from fastapi import HTTPException, status
from utils.logger import logger
from config import settings

class ResumeParser:
    def validate_file(self, filename: str, content: bytes):
        """Perform size, extension, and magic-number signature checks."""
        # 1. Check size
        if len(content) > settings.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File exceeds maximum allowed size of {settings.MAX_FILE_SIZE / (1024*1024):.1f}MB."
            )
            
        # 2. Check extension
        ext = filename.split(".")[-1].lower() if "." in filename else ""
        if ext not in settings.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Unsupported file format. Only PDF and DOCX files are allowed."
            )
            
        # 3. Check signature (magic numbers)
        if ext == "pdf":
            if not content.startswith(b"%PDF-"):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Invalid PDF file. The file headers do not match a PDF signature."
                )
        elif ext == "docx":
            # DOCX is essentially a ZIP archive
            if not content.startswith(b"PK\x03\x04"):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Invalid DOCX file. The file headers do not match a DOCX (ZIP) signature."
                )

    def parse_pdf(self, content: bytes) -> str:
        """Parse text content from PDF bytes using pdfplumber."""
        try:
            text_content = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                if not pdf.pages:
                    raise ValueError("The PDF file contains no pages.")
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            full_text = "\n".join(text_content).strip()
            if not full_text:
                raise ValueError("The PDF is scanned or does not contain extractable text.")
            return full_text
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            if isinstance(e, HTTPException):
                raise e
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to parse PDF: {str(e)}"
            )

    def parse_docx(self, content: bytes) -> str:
        """Parse text content from DOCX bytes using python-docx."""
        try:
            doc = docx.Document(io.BytesIO(content))
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_content.append(paragraph.text)
                    
            # Extract from tables as well
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_content.append(cell.text)
                            
            full_text = "\n".join(text_content).strip()
            if not full_text:
                raise ValueError("The DOCX file does not contain extractable text.")
            return full_text
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            if isinstance(e, HTTPException):
                raise e
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to parse DOCX: {str(e)}"
            )

    def parse_resume(self, filename: str, content: bytes) -> str:
        """Entrypoint for validating and parsing resume files."""
        self.validate_file(filename, content)
        ext = filename.split(".")[-1].lower()
        if ext == "pdf":
            return self.parse_pdf(content)
        elif ext == "docx":
            return self.parse_docx(content)
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Unsupported extension"
            )
